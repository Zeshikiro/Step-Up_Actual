# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

title = doc.add_heading('3.X Software Development Life Cycle: Agile Methodology', level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

p1 = doc.add_paragraph('The development of the ')
p1.add_run('STEP-UP').bold = True
p1.add_run(' mobile application was executed using an ')
p1.add_run('Agile Software Development Life Cycle (SDLC)').bold = True
p1.add_run('. This methodology was chosen over a traditional Waterfall model to accommodate the complex integration of emerging technologies (AR, GPS tracking, physical sensors) and to allow for rapid, iterative prototyping. The project\\'s timeline and development logs demonstrate a strict adherence to core Agile principles: iterative sprint cycles, continuous adaptation to changing requirements, vertical slicing of features, and daily progress tracking.')

doc.add_heading('1. Sprint-Based Iterative Development', level=2)
doc.add_paragraph('The project timeline was broken down into seven distinct, highly focused development phases, each acting as an Agile Sprint. These sprints ranged from one to three weeks and were designed to deliver shippable, functional increments of the application:')
sprints = [
    ('Sprint 1 (Phase 1: UI & Routing):', ' Established the core navigational foundation (April 11 - April 24).'),
    ('Sprint 2 (Phase 2: Security & Backend):', ' Integrated secure session management and authentication (May 4 - May 14).'),
    ('Sprint 3 (Phase 3: Hardware & Gamify):', ' Connected the native Android pedometer to the UI logic (May 12 - May 22).'),
    ('Sprint 4 (Phase 4: Avatar System):', ' A concentrated one-week sprint to rapidly build and finalize the 3D full-body mesh customization system (May 26 - June 1).'),
    ('Sprint 5 (Phase 5: Mapbox & Cloud):', ' Integrated complex GPS APIs and Firebase cloud saves (May 31 - June 16).'),
    ('Sprint 6 (Phase 6: Web & Leaderboard):', ' Developed the promotional React Web App and live ranking system (May 20 - June 4).'),
    ('Sprint 7 (Phase 7: Polish & Compliance):', ' Addressed major Android platform compliance, UI/UX polish, and anti-cheat enforcement (June 21 - August 3).')
]
for title_str, text_str in sprints:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title_str).bold = True
    p.add_run(text_str)

doc.add_heading('2. Adaptability and Rapid Pivoting', level=2)
doc.add_paragraph('A defining characteristic of Agile methodology is the prioritization of responding to change over following a rigid plan. Throughout the development of STEP-UP, technical roadblocks necessitated immediate architectural pivots, which were seamlessly absorbed into the workflow:')
pivots = [
    ('Camera & UX Pivot (April 21):', ' The initial 3D Map camera approach was discarded in favor of a dedicated AR Scene to drastically improve user immersion and usability.'),
    ('Architecture Refactor (May 26):', ' The fragmented body-part avatar system was proving destructive to the rig\\'s internal hierarchy. The team rapidly pivoted to a lightweight, full-body mesh-swapping system, avoiding weeks of technical debt.'),
    ('API Migration (June 13):', ' When the Google Places API introduced restrictive credit card authorization limits, the team immediately migrated to the Yelp Fusion API for POI data generation. Under a Waterfall model, this would have caused catastrophic delays, but Agile allowed for instant adaptation.')
]
for title_str, text_str in pivots:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title_str).bold = True
    p.add_run(text_str)

doc.add_heading('3. Vertical Slicing and Continuous Integration', level=2)
doc.add_paragraph('Instead of building the database, followed by the UI, and then the gameplay logic in isolated, horizontal layers, STEP-UP was built via vertical slices. Working software was prioritized at every stage:')
slices = [
    ('By May 11,', ' a fully functional loop was already operating: Login -> EULA -> BMI -> Settings.'),
    ('By May 12,', ' that loop was immediately wired to the physical hardware step counter.'),
    ('By continuously integrating front-end UI', ' with back-end hardware and Firebase databases, the team ensured the application was testable and playable at the end of nearly every development day.')
]
for title_str, text_str in slices:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title_str).bold = True
    p.add_run(text_str)

doc.add_heading('4. Daily Stand-up and Task Tracking', level=2)
doc.add_paragraph('The project maintained a comprehensive, chronological Development Log that mirrored the format of Agile Daily Stand-ups. Progress tracking consistently evaluated three key metrics:')
doc.add_paragraph('1. What was accomplished today.', style='List Number')
doc.add_paragraph('2. Challenges addressed and bugs squashed.', style='List Number')
doc.add_paragraph('3. Tomorrow\\'s agenda and backlog priority.', style='List Number')
doc.add_paragraph('This systematic documentation ensured that technical debt was addressed instantly (e.g., resolving Unity compiler errors and Android permission conflicts within 24 hours of discovery) and maintained clear, actionable goals for the following sprint.')

doc.save('D:/Unity Projects/Step - Up/Capstone-Context/Agile_Methodology_Defense.docx')
